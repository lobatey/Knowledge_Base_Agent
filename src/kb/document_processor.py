import os
import re
import io
import zipfile
import traceback
from pathlib import Path
from collections import Counter
from typing import Any, Iterable

import chardet
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

try:
    import pdfplumber
except Exception:
    pdfplumber = None

try:
    import fitz  # PyMuPDF
except Exception:
    fitz = None

try:
    import pytesseract
    from PIL import Image
except Exception:
    pytesseract = None
    Image = None

try:
    from docx import Document as DocxDocument
except Exception:
    DocxDocument = None


SUPPORTED_EXTENSIONS = {
    ".pdf": "pdf",
    ".docx": "docx",
    ".txt": "txt",
    ".md": "markdown",
}

OCR_LANG = os.getenv("OCR_LANG", "chi_sim+eng")
ENABLE_OCR = os.getenv("ENABLE_OCR", "true").lower() in {"1", "true", "yes", "y"}

DEFAULT_CHUNK_SIZE = int(os.getenv("KB_CHUNK_SIZE", "900"))
DEFAULT_CHUNK_OVERLAP = int(os.getenv("KB_CHUNK_OVERLAP", "120"))


def document_type(file_path: Path) -> str:
    suffix = file_path.suffix.lower()

    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"不支持的文件类型，仅支持: {', '.join(SUPPORTED_EXTENSIONS.keys())}"
        )

    return SUPPORTED_EXTENSIONS[suffix]


def normalize_text(text: str) -> str:
    """
    基础文本清洗：
    1. 去掉不可见控制字符
    2. 统一空格
    3. 合并异常空行
    4. 删除常见页码行
    5. 删除明显乱码残留
    """

    if not text:
        return ""

    text = text.replace("\ufeff", "")
    text = text.replace("\u00a0", " ")
    text = text.replace("\u200b", "")
    text = text.replace("\u200c", "")
    text = text.replace("\u200d", "")

    # 删除大部分控制字符，但保留换行和制表符
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", text)

    # 常见页码
    page_number_patterns = [
        r"(?m)^\s*\d+\s*$",
        r"(?m)^\s*第\s*\d+\s*页\s*$",
        r"(?m)^\s*第\s*\d+\s*页\s*/\s*共\s*\d+\s*页\s*$",
        r"(?m)^\s*Page\s+\d+(\s+of\s+\d+)?\s*$",
        r"(?m)^\s*-\s*\d+\s*-\s*$",
    ]

    for pattern in page_number_patterns:
        text = re.sub(pattern, "", text, flags=re.IGNORECASE)

    # 空格和空行
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n[ \t]+", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)

    return text.strip()


def is_probably_formula(line: str) -> bool:
    """
    粗略判断一行是否像公式。
    目标不是做数学理解，而是防止公式被当成普通句子切碎。
    """

    if not line:
        return False

    line = line.strip()

    formula_chars = set("=+-*/^_∑∫√≈≠≤≥∞αβγδλμσθπΩ∆∂")
    formula_score = sum(1 for char in line if char in formula_chars)

    has_latex = bool(
        re.search(
            r"(\\frac|\\sum|\\int|\\sqrt|\\begin|\\end|\$.*\$)",
            line,
        )
    )

    has_math_pattern = bool(
        re.search(
            r"[a-zA-Z]\s*=\s*[^，。；;]+",
            line,
        )
    )

    return has_latex or has_math_pattern or formula_score >= 3

CAPTION_PATTERNS = [
    # 中文图注：图1、图 1、图1-2、图 1.2、图一
    (
        "figure",
        re.compile(
            r"^\s*(图|图表|Figure|Fig\.?)\s*"
            r"([0-9一二三四五六七八九十百]+([\-\.．—_][0-9一二三四五六七八九十百]+)*)"
            r"[\s:：、.\-—]*"
            r"(.+)?\s*$",
            re.IGNORECASE,
        ),
    ),
    # 中文表题：表1、表 1、表1-2、表 1.2、表一
    (
        "table",
        re.compile(
            r"^\s*(表|Table)\s*"
            r"([0-9一二三四五六七八九十百]+([\-\.．—_][0-9一二三四五六七八九十百]+)*)"
            r"[\s:：、.\-—]*"
            r"(.+)?\s*$",
            re.IGNORECASE,
        ),
    ),
    # 公式题注：式（1）、公式（1）、Equation 1
    (
        "formula",
        re.compile(
            r"^\s*(式|公式|Equation|Eq\.?)\s*"
            r"[（(]?\s*([0-9一二三四五六七八九十百]+([\-\.．—_][0-9一二三四五六七八九十百]+)*)\s*[）)]?"
            r"[\s:：、.\-—]*"
            r"(.+)?\s*$",
            re.IGNORECASE,
        ),
    ),
]


def parse_caption_line(line: str) -> dict[str, str] | None:
    """
    判断一行是否是图注、表题或公式题注。

    返回示例：
    {
        "caption_type": "table",
        "caption_label": "表 1",
        "caption_text": "表 1 系统接口说明"
    }
    """

    line = normalize_text(line)

    if not line:
        return None

    # 过长的一行通常不是题注，而是普通正文
    if len(line) > 160:
        return None

    for caption_type, pattern in CAPTION_PATTERNS:
        match = pattern.match(line)

        if not match:
            continue

        prefix = match.group(1)
        number = match.group(2)

        caption_label = f"{prefix} {number}".strip()

        return {
            "caption_type": caption_type,
            "caption_label": caption_label,
            "caption_text": line,
        }

    return None


def extract_caption_docs_from_text_docs(documents: list[Document]) -> list[Document]:
    """
    从普通 text 文档中识别图注、表题、公式题注。

    做法：
    1. 如果某一行像 “图 1 xxx” 或 “表 1 xxx”，单独抽成 caption doc
    2. 剩下的正文仍然保留为 text doc

    注意：
    这里不是删除题注，而是把题注从普通正文里拆出来，方便后面绑定。
    """

    result: list[Document] = []

    for doc in documents:
        if doc.metadata.get("content_type") != "text":
            result.append(doc)
            continue

        text_buffer: list[str] = []

        def flush_text_buffer():
            nonlocal text_buffer

            content = normalize_text("\n".join(text_buffer))
            text_buffer = []

            if content:
                result.append(
                    Document(
                        page_content=content,
                        metadata=dict(doc.metadata),
                    )
                )

        for line in doc.page_content.splitlines():
            stripped = normalize_text(line)
            caption_info = parse_caption_line(stripped)

            if caption_info:
                flush_text_buffer()

                caption_metadata = dict(doc.metadata)
                caption_metadata["content_type"] = "caption"
                caption_metadata.update(caption_info)

                result.append(
                    Document(
                        page_content=f"[题注]\n{caption_info['caption_text']}",
                        metadata=caption_metadata,
                    )
                )
            else:
                text_buffer.append(line)

        flush_text_buffer()

    return result

def attach_captions_to_structured_docs(documents: list[Document]) -> list[Document]:
    """
    将 caption doc 绑定到邻近的 table / image_ocr / formula / ocr_text 文档上。

    绑定规则：
    1. table caption 优先绑定到最近的 table
    2. figure caption 优先绑定到最近的 image_ocr 或 ocr_text
    3. formula caption 优先绑定到最近的 formula
    4. 优先考虑同一页
    5. 支持题注在结构块前面，也支持题注在结构块后面

    绑定后：
    - caption 会写入结构块 metadata
    - caption 会拼接进结构块 page_content
    - caption doc 本身默认不再单独保留，避免重复入库
    """

    if not documents:
        return []

    docs = [
        Document(
            page_content=doc.page_content,
            metadata=dict(doc.metadata),
        )
        for doc in documents
    ]

    used_caption_indexes: set[int] = set()

    target_types_by_caption_type = {
        "table": {"table"},
        "figure": {"image_ocr", "ocr_text", "image_caption"},
        "formula": {"formula"},
    }

    def same_page(a: Document, b: Document) -> bool:
        page_a = a.metadata.get("page")
        page_b = b.metadata.get("page")

        if page_a is None or page_b is None:
            return True

        return page_a == page_b

    def find_best_target(caption_index: int, caption_doc: Document) -> int | None:
        caption_type = caption_doc.metadata.get("caption_type")
        target_types = target_types_by_caption_type.get(caption_type)

        if not target_types:
            return None

        best_index = None
        best_score = 10**9

        for index, candidate in enumerate(docs):
            if index == caption_index:
                continue

            candidate_type = candidate.metadata.get("content_type")

            if candidate_type not in target_types:
                continue

            if not same_page(caption_doc, candidate):
                continue

            distance = abs(index - caption_index)

            # 距离太远就不绑定，避免误绑
            if distance > 3:
                continue

            # 方向优先级：
            # 表题通常在表格前面，所以 caption_index < table_index 更优
            # 图注通常在图片后面，所以 caption_index > image_index 更优
            # 公式题注两边都可能出现，所以不强偏向
            direction_penalty = 0

            if caption_type == "table":
                if caption_index > index:
                    direction_penalty = 1

            elif caption_type == "figure":
                if caption_index < index:
                    direction_penalty = 1

            elif caption_type == "formula":
                direction_penalty = 0

            score = distance * 10 + direction_penalty

            if score < best_score:
                best_score = score
                best_index = index

        return best_index

    for index, doc in enumerate(docs):
        if doc.metadata.get("content_type") != "caption":
            continue

        target_index = find_best_target(index, doc)

        if target_index is None:
            continue

        target_doc = docs[target_index]

        caption_text = doc.metadata.get("caption_text") or doc.page_content
        caption_type = doc.metadata.get("caption_type")
        caption_label = doc.metadata.get("caption_label")

        caption_text = normalize_text(caption_text)

        if not caption_text:
            continue

        target_doc.metadata["caption"] = caption_text
        target_doc.metadata["caption_type"] = caption_type
        target_doc.metadata["caption_label"] = caption_label

        target_content_type = target_doc.metadata.get("content_type")

        if target_content_type == "table":
            prefix = "[表格题注]"
        elif target_content_type in {"image_ocr", "ocr_text", "image_caption"}:
            prefix = "[图片题注]"
        elif target_content_type == "formula":
            prefix = "[公式题注]"
        else:
            prefix = "[题注]"

        # 避免重复拼接题注
        if caption_text not in target_doc.page_content:
            target_doc.page_content = normalize_text(
                f"{prefix}\n{caption_text}\n\n{target_doc.page_content}"
            )

        used_caption_indexes.add(index)

    result: list[Document] = []

    for index, doc in enumerate(docs):
        content_type = doc.metadata.get("content_type")

        # 已经成功绑定到表格/图片/公式的 caption，不再单独入库
        if content_type == "caption" and index in used_caption_indexes:
            continue

        # 没绑定成功的 caption 仍然保留，避免信息丢失
        result.append(doc)

    return result


def table_to_markdown(table: list[list[Any]]) -> str:
    """
    把 pdfplumber / python-docx 提取到的二维表格转成 Markdown 表格。
    这样入库后，行列关系比普通文本更稳定。
    """

    if not table:
        return ""

    cleaned_rows = []

    for row in table:
        if not row:
            continue

        cleaned_row = []
        for cell in row:
            value = "" if cell is None else str(cell)
            value = normalize_text(value)
            value = value.replace("\n", "<br>")
            cleaned_row.append(value)

        if any(cell.strip() for cell in cleaned_row):
            cleaned_rows.append(cleaned_row)

    if not cleaned_rows:
        return ""

    max_cols = max(len(row) for row in cleaned_rows)

    normalized_rows = []
    for row in cleaned_rows:
        row = row + [""] * (max_cols - len(row))
        normalized_rows.append(row)

    header = normalized_rows[0]
    body = normalized_rows[1:] if len(normalized_rows) > 1 else []

    markdown_lines = []
    markdown_lines.append("| " + " | ".join(header) + " |")
    markdown_lines.append("| " + " | ".join(["---"] * max_cols) + " |")

    for row in body:
        markdown_lines.append("| " + " | ".join(row) + " |")

    return "\n".join(markdown_lines).strip()


def make_doc(
    content: str,
    source: str,
    file_type: str,
    content_type: str,
    page: int | None = None,
    extra_metadata: dict[str, Any] | None = None,
) -> Document:
    content = normalize_text(content)

    metadata: dict[str, Any] = {
        "source": source,
        "file_type": file_type,
        "content_type": content_type,
    }

    if page is not None:
        metadata["page"] = page

    if extra_metadata:
        metadata.update(extra_metadata)

    return Document(page_content=content, metadata=metadata)


def remove_repeated_lines(documents: list[Document], min_repeat: int = 3) -> list[Document]:
    """
    删除跨页重复出现的短行。
    主要用于去除页眉、页脚、水印、公司名称、保密标识等。
    """

    line_counter: Counter[str] = Counter()

    for doc in documents:
        if doc.metadata.get("content_type") != "text":
            continue

        lines = {
            line.strip()
            for line in doc.page_content.splitlines()
            if line.strip()
        }
        line_counter.update(lines)

    repeated_lines = {
        line
        for line, count in line_counter.items()
        if count >= min_repeat and len(line) <= 100
    }

    cleaned_docs: list[Document] = []

    for doc in documents:
        if doc.metadata.get("content_type") != "text":
            cleaned_docs.append(doc)
            continue

        lines = []
        for line in doc.page_content.splitlines():
            stripped = line.strip()
            if stripped in repeated_lines:
                continue
            lines.append(line)

        cleaned_content = normalize_text("\n".join(lines))

        if cleaned_content:
            cleaned_docs.append(
                Document(
                    page_content=cleaned_content,
                    metadata=doc.metadata,
                )
            )

    return cleaned_docs


def split_formula_blocks(documents: list[Document]) -> list[Document]:
    """
    把普通文本里的公式行单独拆出来。
    这样公式不会被普通段落切分器随意切碎。
    """

    result: list[Document] = []

    for doc in documents:
        if doc.metadata.get("content_type") != "text":
            result.append(doc)
            continue

        text_lines: list[str] = []
        formula_index = 0

        for line in doc.page_content.splitlines():
            stripped = line.strip()

            if is_probably_formula(stripped):
                if text_lines:
                    content = normalize_text("\n".join(text_lines))
                    if content:
                        result.append(Document(page_content=content, metadata=doc.metadata))
                    text_lines = []

                formula_metadata = dict(doc.metadata)
                formula_metadata["content_type"] = "formula"
                formula_metadata["formula_index"] = formula_index

                result.append(
                    Document(
                        page_content=f"[公式]\n{stripped}",
                        metadata=formula_metadata,
                    )
                )
                formula_index += 1
            else:
                text_lines.append(line)

        if text_lines:
            content = normalize_text("\n".join(text_lines))
            if content:
                result.append(Document(page_content=content, metadata=doc.metadata))

    return result


def ocr_image_bytes(image_bytes: bytes) -> str:
    """
    对图片字节做 OCR。
    如果没有安装 pytesseract / pillow / tesseract，则返回空字符串。
    """

    if not ENABLE_OCR or pytesseract is None or Image is None:
        return ""

    try:
        image = Image.open(io.BytesIO(image_bytes)).convert("RGB")
        text = pytesseract.image_to_string(image, lang=OCR_LANG)
        return normalize_text(text)
    except Exception:
        return ""


def ocr_pdf_page(file_path: Path, page_index: int, zoom: float = 2.0) -> str:
    """
    对扫描版 PDF 的某一页做 OCR。
    page_index 从 0 开始。
    """

    if not ENABLE_OCR or fitz is None or pytesseract is None or Image is None:
        return ""

    try:
        pdf = fitz.open(str(file_path))
        page = pdf.load_page(page_index)

        matrix = fitz.Matrix(zoom, zoom)
        pix = page.get_pixmap(matrix=matrix, alpha=False)

        image = Image.open(io.BytesIO(pix.tobytes("png"))).convert("RGB")
        text = pytesseract.image_to_string(image, lang=OCR_LANG)

        pdf.close()

        return normalize_text(text)
    except Exception:
        traceback.print_exc()
        return ""


def extract_pdf_documents(file_path: Path) -> list[Document]:
    """
    PDF 解析：
    1. 用 pdfplumber 提取每页正文
    2. 提取表格，并转成 Markdown
    3. 如果某页没有文字，尝试 OCR
    """

    if pdfplumber is None:
        raise RuntimeError("缺少 pdfplumber，请先执行 pip install pdfplumber")

    docs: list[Document] = []
    file_type = "pdf"

    with pdfplumber.open(str(file_path)) as pdf:
        for page_index, page in enumerate(pdf.pages, start=1):
            text = page.extract_text(x_tolerance=1, y_tolerance=3) or ""
            text = normalize_text(text)

            if text:
                docs.append(
                    make_doc(
                        content=text,
                        source=file_path.name,
                        file_type=file_type,
                        content_type="text",
                        page=page_index,
                    )
                )
            else:
                ocr_text = ocr_pdf_page(file_path, page_index - 1)
                if ocr_text:
                    docs.append(
                        make_doc(
                            content=f"[OCR识别文本]\n{ocr_text}",
                            source=file_path.name,
                            file_type=file_type,
                            content_type="ocr_text",
                            page=page_index,
                        )
                    )

            try:
                tables = page.extract_tables() or []
            except Exception:
                tables = []

            for table_index, table in enumerate(tables):
                markdown_table = table_to_markdown(table)

                if markdown_table:
                    docs.append(
                        make_doc(
                            content=f"[表格]\n{markdown_table}",
                            source=file_path.name,
                            file_type=file_type,
                            content_type="table",
                            page=page_index,
                            extra_metadata={
                                "table_index": table_index,
                            },
                        )
                    )

    return docs


def iter_docx_blocks(docx_document) -> Iterable[tuple[str, Any]]:
    """
    按 Word 正文顺序读取段落和表格。
    python-docx 默认 Document.paragraphs 和 Document.tables 会分开读，
    这样容易丢失顺序，所以这里按 XML body 顺序遍历。
    """

    from docx.table import Table
    from docx.text.paragraph import Paragraph

    body = docx_document.element.body

    for child in body.iterchildren():
        if child.tag.endswith("}p"):
            yield "paragraph", Paragraph(child, docx_document)
        elif child.tag.endswith("}tbl"):
            yield "table", Table(child, docx_document)


def extract_docx_formulas(file_path: Path) -> list[str]:
    """
    尝试从 DOCX XML 里提取 Word 公式文本。
    这不是完整的公式识别，但可以尽量避免公式完全丢失。
    """

    formulas: list[str] = []

    try:
        with zipfile.ZipFile(file_path) as zf:
            xml_names = [
                name
                for name in zf.namelist()
                if name.startswith("word/") and name.endswith(".xml")
            ]

            for xml_name in xml_names:
                xml_text = zf.read(xml_name).decode("utf-8", errors="ignore")

                # 粗略捕获 oMath / oMathPara 片段
                blocks = re.findall(
                    r"<m:oMathPara[\s\S]*?</m:oMathPara>|<m:oMath[\s\S]*?</m:oMath>",
                    xml_text,
                )

                for block in blocks:
                    texts = re.findall(r"<(?:m:t|w:t)[^>]*>(.*?)</(?:m:t|w:t)>", block)
                    formula = "".join(texts)
                    formula = normalize_text(formula)

                    if formula:
                        formulas.append(formula)

    except Exception:
        pass

    return formulas


def extract_docx_image_ocr(file_path: Path) -> list[str]:
    """
    从 DOCX 的 word/media 目录提取图片，并对图片做 OCR。
    """

    image_texts: list[str] = []

    if not ENABLE_OCR:
        return image_texts

    try:
        with zipfile.ZipFile(file_path) as zf:
            image_names = [
                name
                for name in zf.namelist()
                if name.startswith("word/media/")
                and name.lower().endswith((".png", ".jpg", ".jpeg", ".bmp", ".tif", ".tiff"))
            ]

            for image_name in image_names:
                image_bytes = zf.read(image_name)
                text = ocr_image_bytes(image_bytes)

                if text:
                    image_texts.append(text)

    except Exception:
        pass

    return image_texts


def extract_docx_documents(file_path: Path) -> list[Document]:
    """
    DOCX 解析：
    1. 按正文顺序解析段落和表格
    2. 表格转 Markdown
    3. 尝试提取 Word 公式
    4. 尝试对文档图片 OCR
    """

    if DocxDocument is None:
        raise RuntimeError("缺少 python-docx，请先执行 pip install python-docx")

    docs: list[Document] = []
    file_type = "docx"
    docx_document = DocxDocument(str(file_path))

    paragraph_buffer: list[str] = []
    table_index = 0

    def flush_paragraph_buffer():
        nonlocal paragraph_buffer

        content = normalize_text("\n".join(paragraph_buffer))
        paragraph_buffer = []

        if content:
            docs.append(
                make_doc(
                    content=content,
                    source=file_path.name,
                    file_type=file_type,
                    content_type="text",
                )
            )

    for block_type, block in iter_docx_blocks(docx_document):
        if block_type == "paragraph":
            text = normalize_text(block.text)
            if text:
                paragraph_buffer.append(text)

        elif block_type == "table":
            flush_paragraph_buffer()

            table_data: list[list[str]] = []
            for row in block.rows:
                table_data.append([cell.text for cell in row.cells])

            markdown_table = table_to_markdown(table_data)

            if markdown_table:
                docs.append(
                    make_doc(
                        content=f"[表格]\n{markdown_table}",
                        source=file_path.name,
                        file_type=file_type,
                        content_type="table",
                        extra_metadata={
                            "table_index": table_index,
                        },
                    )
                )
                table_index += 1

    flush_paragraph_buffer()

    for formula_index, formula in enumerate(extract_docx_formulas(file_path)):
        docs.append(
            make_doc(
                content=f"[公式]\n{formula}",
                source=file_path.name,
                file_type=file_type,
                content_type="formula",
                extra_metadata={
                    "formula_index": formula_index,
                },
            )
        )

    for image_index, image_text in enumerate(extract_docx_image_ocr(file_path)):
        docs.append(
            make_doc(
                content=f"[图片OCR文本]\n{image_text}",
                source=file_path.name,
                file_type=file_type,
                content_type="image_ocr",
                extra_metadata={
                    "image_index": image_index,
                },
            )
        )

    return docs


def read_text_file_with_fallback(file_path: Path) -> str:
    """
    TXT / MD 编码兼容：
    优先 UTF-8，失败后用 chardet 猜测编码。
    """

    raw = file_path.read_bytes()

    try:
        return raw.decode("utf-8")
    except UnicodeDecodeError:
        detected = chardet.detect(raw)
        encoding = detected.get("encoding") or "utf-8"
        return raw.decode(encoding, errors="ignore")


def extract_plain_text_documents(file_path: Path) -> list[Document]:
    file_type = document_type(file_path)
    text = read_text_file_with_fallback(file_path)
    text = normalize_text(text)

    if not text:
        return []

    return [
        make_doc(
            content=text,
            source=file_path.name,
            file_type=file_type,
            content_type="text",
        )
    ]


def extract_documents(file_path: Path) -> list[Document]:
    """
    对外统一入口：
    根据文件类型调用不同解析器。

    处理顺序：
    1. 提取 PDF / DOCX / TXT / MD 内容
    2. 删除空内容
    3. 去掉重复页眉页脚
    4. 从正文中拆出图注、表题、公式题注
    5. 识别并拆分公式
    6. 将题注绑定到邻近的表格、图片 OCR、公式 chunk
    """

    suffix = file_path.suffix.lower()

    if suffix == ".pdf":
        docs = extract_pdf_documents(file_path)
    elif suffix == ".docx":
        docs = extract_docx_documents(file_path)
    elif suffix in {".txt", ".md"}:
        docs = extract_plain_text_documents(file_path)
    else:
        raise ValueError(f"不支持的文件类型: {suffix}")

    docs = [doc for doc in docs if normalize_text(doc.page_content)]

    # 1. 去页眉页脚
    docs = remove_repeated_lines(docs)

    # 2. 先把题注从普通正文中拆出来
    docs = extract_caption_docs_from_text_docs(docs)

    # 3. 再把公式从正文中拆出来
    docs = split_formula_blocks(docs)

    # 4. 最后把题注绑定到表格、图片、公式
    docs = attach_captions_to_structured_docs(docs)

    return docs



def smart_split_documents(documents: list[Document]) -> list[Document]:
    """
    智能切分策略：

    1. 表格、公式、图片 OCR 文本：尽量作为独立块保留
    2. 普通正文：按标题、段落、句号等优先级切分
    3. 每个 chunk 保留 source / page / content_type 等元数据
    """

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=DEFAULT_CHUNK_SIZE,
        chunk_overlap=DEFAULT_CHUNK_OVERLAP,
        separators=[
            "\n# ",
            "\n## ",
            "\n### ",
            "\n\n",
            "\n",
            "。", "！", "？",
            ". ", "! ", "? ",
            "；", ";",
            "，", ",",
            " ",
            "",
        ],
    )

    final_chunks: list[Document] = []

    for doc in documents:
        content_type = doc.metadata.get("content_type", "text")
        content = normalize_text(doc.page_content)

        if not content:
            continue

        # 这些结构化内容不希望被轻易切碎
        if content_type in {
            "table",
            "formula",
            "image_ocr",
            "ocr_text",
        }:
            if len(content) <= DEFAULT_CHUNK_SIZE * 1.5:
                final_chunks.append(
                    Document(
                        page_content=content,
                        metadata=dict(doc.metadata),
                    )
                )
            else:
                # 如果特别长，仍然切分，但保留 content_type
                final_chunks.extend(splitter.split_documents([doc]))

        else:
            final_chunks.extend(splitter.split_documents([doc]))

    for index, chunk in enumerate(final_chunks):
        chunk.metadata["chunk_index"] = index

    return final_chunks
